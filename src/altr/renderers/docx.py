from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .. import theme
from ..markdown import markdown_to_blocks, split_inline
from ..schemas import (
    BulletList,
    DocumentSpec,
    Heading,
    Image,
    Markdown,
    NumberedList,
    PageBreak,
    Paragraph,
    Table,
)

# (size pt, color, space before pt) per heading style in the default theme.
_HEADING_STYLES = {
    "Heading 1": (16, theme.ACCENT_DEEP, 18),
    "Heading 2": (13, theme.ACCENT_DEEP, 14),
    "Heading 3": (11.5, theme.INK, 12),
    "Heading 4": (11, theme.INK_SECONDARY, 10),
}


def render_document(
    spec: DocumentSpec, out_dir: Path, template: Path | None = None
) -> Path:
    from . import output_path

    doc = Document(str(template)) if template else Document()
    styled = template is None  # never restyle a user-provided brand template
    if styled:
        _apply_theme(doc)
    if spec.title:
        para = doc.add_heading(spec.title, 0)
        if styled:
            _add_bottom_rule(para, theme.ACCENT)

    for block in spec.blocks:
        _add_block(doc, block, styled)

    path = output_path(out_dir, spec.filename, ".docx")
    doc.save(str(path))
    return path


def _apply_theme(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"  # pin: unset, the body falls back to a serif
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(theme.INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(theme.ACCENT_DEEP)

    for name, (size, color, space_before) in _HEADING_STYLES.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(4)

    styles["Caption"].font.color.rgb = RGBColor.from_string(theme.INK_SECONDARY)


def _add_block(doc: Document, block, styled: bool) -> None:
    if isinstance(block, Heading):
        doc.add_heading(block.text, block.level)
    elif isinstance(block, Paragraph):
        _add_styled_paragraph(doc, block.text)
    elif isinstance(block, BulletList):
        for item in block.items:
            _add_styled_paragraph(doc, item, style="List Bullet")
    elif isinstance(block, NumberedList):
        for item in block.items:
            _add_styled_paragraph(doc, item, style="List Number")
    elif isinstance(block, Table):
        _add_table(doc, block, styled)
    elif isinstance(block, Image):
        image_path = Path(block.path)
        if not image_path.is_file():
            raise FileNotFoundError(f"image not found: {block.path}")
        width = Inches(block.width_inches) if block.width_inches else None
        doc.add_picture(str(image_path), width=width)
        if block.caption:
            doc.add_paragraph(block.caption, style="Caption")
    elif isinstance(block, Markdown):
        for sub_block in markdown_to_blocks(block.text):
            _add_block(doc, sub_block, styled)
    elif isinstance(block, PageBreak):
        doc.add_page_break()


def _add_table(doc: Document, block: Table, styled: bool) -> None:
    table = doc.add_table(rows=1, cols=len(block.headers))
    if not styled:
        table.style = "Table Grid"
    for i, header in enumerate(block.headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        if styled:
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            _shade_cell(cell, theme.HEADER_FILL)
    for r, row in enumerate(block.rows):
        cells = table.add_row().cells
        for i, value in enumerate(row[: len(block.headers)]):
            cells[i].text = value
            if styled and r % 2 == 1:
                _shade_cell(cells[i], theme.ROW_BAND)
    if styled:
        _set_table_borders(table)


def _add_styled_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    """Add a paragraph, turning inline **bold**/*italic*/`code` into runs."""
    para = doc.add_paragraph(style=style)
    for kind, part in split_inline(text):
        run = para.add_run(part)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Courier New"


# --- low-level OOXML helpers (python-docx has no API for these) --------------


def _add_bottom_rule(paragraph, color: str, size: int = 12) -> None:
    """Draw an accent rule under a paragraph (used below the document title)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))  # eighths of a point
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _shade_cell(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table) -> None:
    """Horizontal hairlines only - no vertical grid - for a clean, open look."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("insideH", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), theme.GRIDLINE)
        borders.append(el)
    for edge in ("top", "left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)
