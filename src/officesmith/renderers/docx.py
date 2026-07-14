from __future__ import annotations

from pathlib import Path

from docx import Document

from ..schemas import (
    BulletList,
    DocumentSpec,
    Heading,
    NumberedList,
    PageBreak,
    Paragraph,
    Table,
)


def render_document(spec: DocumentSpec, out_dir: Path) -> Path:
    from . import output_path

    doc = Document()
    if spec.title:
        doc.add_heading(spec.title, 0)

    for block in spec.blocks:
        if isinstance(block, Heading):
            doc.add_heading(block.text, block.level)
        elif isinstance(block, Paragraph):
            doc.add_paragraph(block.text)
        elif isinstance(block, BulletList):
            for item in block.items:
                doc.add_paragraph(item, style="List Bullet")
        elif isinstance(block, NumberedList):
            for item in block.items:
                doc.add_paragraph(item, style="List Number")
        elif isinstance(block, Table):
            table = doc.add_table(rows=1, cols=len(block.headers))
            table.style = "Table Grid"
            for i, header in enumerate(block.headers):
                cell = table.rows[0].cells[i]
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
            for row in block.rows:
                cells = table.add_row().cells
                for i, value in enumerate(row[: len(block.headers)]):
                    cells[i].text = value
        elif isinstance(block, PageBreak):
            doc.add_page_break()

    path = output_path(out_dir, spec.filename, ".docx")
    doc.save(str(path))
    return path
